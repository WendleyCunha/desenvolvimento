import streamlit as st
import pandas as pd
import zipfile
import os
from datetime import datetime
from docx import Document
from io import BytesIO
import database as db  # Conexão com seu Firebase

# --- FUNÇÕES DE NÚCLEO ---

def gerar_word_memoria(dados):
    try:
        doc = Document("carta_preenchida.docx")
        for p in doc.paragraphs:
            for k, v in dados.items():
                if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in dados.items():
                            if f"{{{{{k}}}}}" in p.text: p.text = p.text.replace(f"{{{{{k}}}}}", str(v))
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Erro no template: {e}")
        return None

def exibir(user_role):
    st.markdown("""
        <style>
        .card-carta { background-color: white; padding: 20px; border-radius: 15px; border-left: 5px solid #002366; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 15px; }
        .status-badge { padding: 5px 12px; border-radius: 20px; font-size: 12px; font-weight: bold; background-color: #e2e8f0; }
        </style>
    """, unsafe_allow_html=True)

    st.title("📑 Gestão de Cartas de Débito")
    fire = db.inicializar_db()
    
    # --- CARREGAR DADOS ---
    def obter_cartas():
        docs = fire.collection("cartas_rh").stream()
        return [doc.to_dict() for doc in docs]

    cartas = obter_cartas()
    
    # Definição das Abas
    abas_nomes = ["🆕 Nova Carta", "📋 Painel de Controle", "📦 Fechamento de Lote", "✅ Histórico"]
    if user_role in ["ADM", "GERENTE"]:
        abas_nomes.append("⚠️ Reset Sistema")
    
    tabs = st.tabs(abas_nomes)

    # 1. NOVA CARTA
    with tabs[0]:
        with st.form("f_premium", clear_on_submit=True):
            st.subheader("Informações do Lançamento")
            c1, c2, c3 = st.columns(3)
            nome = c1.text_input("Nome do Colaborador").upper()
            cpf = c2.text_input("CPF")
            cod_cli = c3.text_input("Código do Cliente")
            
            v1, v2, v3 = st.columns(3)
            valor = v1.number_input("Valor R$", min_value=0.0)
            loja = v2.text_input("Loja Origem").upper()
            data_c = v3.date_input("Data da Ocorrência")
            
            motivo = st.text_area("Motivo Detalhado").upper()
            
            if st.form_submit_button("✨ Gerar e Registrar"):
                if nome and cpf and cod_cli:
                    id_carta = datetime.now().strftime("%Y%m%d%H%M%S")
                    dados_fb = {
                        "id": id_carta, "NOME": nome, "CPF": cpf, "COD_CLI": cod_cli, 
                        "VALOR": valor, "LOJA": loja, "DATA": data_c.strftime("%d/%m/%Y"), 
                        "MOTIVO": motivo, "status": "Aguardando Assinatura", "anexo_bin": None,
                        "data_criacao": datetime.now().strftime("%d/%m/%Y %H:%M"),
                        "qtd_parcial": 0 # Campo novo para controle
                    }
                    fire.collection("cartas_rh").document(id_carta).set(dados_fb)
                    st.success("Carta registrada com sucesso!")
                    st.rerun()

    # 2. PAINEL DE CONTROLE (ESTEIRA COM BUSCA)
    with tabs[1]:
        st.subheader("🚀 Esteira de Tratativas")
        busca_e = st.text_input("🔍 Buscar na Esteira (Código Cliente ou Nome)", placeholder="Digite para filtrar...")
        
        lista_painel = [c for c in cartas if c.get('status') == "Aguardando Assinatura"]
        
        if busca_e:
            lista_painel = [c for c in lista_painel if busca_e.upper() in c['NOME'] or busca_e in c['COD_CLI']]

        if not lista_painel:
            st.info("Nenhum item pendente na esteira.")
        else:
            for c in lista_painel:
                with st.container():
                    st.markdown(f'''
                        <div class="card-carta">
                            <strong>{c["NOME"]}</strong> | Lo_ja: {c["LOJA"]} | Cliente: {c["COD_CLI"]}<br>
                            <small>Motivo: {c["MOTIVO"]}</small>
                        </div>
                    ''', unsafe_allow_html=True)
                    
                    # ÁREA DE TRATATIVA
                    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
                    
                    # Opção 1: Compra Total
                    if col1.button("✅ TOTAL", key=f"tot_{c['id']}", use_container_width=True):
                        fire.collection("cartas_rh").document(c['id']).update({"status": "Tratado: Total"})
                        st.rerun()
                    
                    # Opção 2: Compra Parcial
                    with col2:
                        with st.popover("➕ PARCIAL", use_container_width=True):
                            qtdp = st.number_input("Qtd Comprada:", min_value=1, key=f"q_{c['id']}")
                            if st.button("Confirmar Parcial", key=f"btp_{c['id']}"):
                                fire.collection("cartas_rh").document(c['id']).update({
                                    "status": "Tratado: Parcial",
                                    "qtd_parcial": qtdp
                                })
                                st.rerun()
                    
                    # Opção 3: Sem Encomenda
                    if col3.button("🔍 SEM ENC.", key=f"sem_{c['id']}", use_container_width=True):
                        fire.collection("cartas_rh").document(c['id']).update({"status": "Tratado: Sem Encomenda"})
                        st.rerun()

                    # Opção 4: Download Word (Para assinar se necessário)
                    dados_w = {"NOME_COLAB": c['NOME'], "CPF": c['CPF'], "CODIGO_CLIENTE": c['COD_CLI'], "VALOR_DEBITO": f"R$ {c['VALOR']:,.2f}", "LOJA_ORIGEM": c['LOJA'], "DATA_COMPRA": c['DATA'], "DESC_DEBITO": c['MOTIVO'], "DATA_LOCAL": f"São Paulo, {datetime.now().strftime('%d/%m/%Y')}"}
                    w_bytes = gerar_word_memoria(dados_w)
                    col4.download_button("📂 DOCX", w_bytes, file_name=f"Carta_{c['NOME']}.docx", key=f"w_{c['id']}", use_container_width=True)
                    
                    # Upload para mudar status para recebida (Finalizar)
                    up = st.file_uploader("Upload da Carta Assinada (Opcional para concluir)", key=f"up_{c['id']}")
                    if up:
                        fire.collection("cartas_rh").document(c['id']).update({
                            "status": "CARTA RECEBIDA", 
                            "anexo_bin": up.getvalue(), 
                            "nome_arquivo": up.name
                        })
                        st.rerun()
                    st.divider()

    # 3. FECHAMENTO DE LOTE
    with tabs[2]:
        # Consideramos tratadas ou recebidas
        prontas = [c for c in cartas if "Tratado" in c.get('status', '') or c.get('status') == "CARTA RECEBIDA"]
        
        if not prontas:
            st.info("Nenhuma carta tratada pronta para fechamento.")
        else:
            st.subheader(f"📦 Lote pronto com {len(prontas)} itens")
            df_preview = pd.DataFrame(prontas)[['NOME', 'COD_CLI', 'VALOR', 'status', 'qtd_parcial']]
            st.dataframe(df_preview, use_container_width=True)
            
            if st.button("🚀 FINALIZAR LOTE E ENVIAR AO HISTÓRICO", type="primary"):
                id_lote = datetime.now().strftime("%Y%m%d_%H%M")
                ids_componentes = [c['id'] for c in prontas]
                
                fire.collection("lotes_rh").document(id_lote).set({
                    "id": id_lote, "data": datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "total": len(prontas), "ids_cartas": ids_componentes
                })
                for c in prontas:
                    fire.collection("cartas_rh").document(c['id']).update({"status": "LOTE_FECHADO"})
                st.success("Lote finalizado!")
                st.rerun()

    # 4. HISTÓRICO
    with tabs[3]:
        col_h1, col_h2 = st.columns([2, 1])
        busca_h = col_h1.text_input("🔎 Pesquisar no Histórico (Nome ou Código)")
        
        # Botão Geral de Exportação
        if cartas:
            df_total = pd.DataFrame(cartas)
            out_geral = BytesIO()
            df_total.to_excel(out_geral, index=False)
            col_h2.download_button("📊 Exportar Tudo (Excel)", out_geral.getvalue(), "Relatorio_Geral_Cartas.xlsx", use_container_width=True)

        docs_lotes = fire.collection("lotes_rh").stream()
        lotes = [d.to_dict() for d in docs_lotes]
        
        for l in sorted(lotes, key=lambda x: x['id'], reverse=True):
            ids_lote = l.get('ids_cartas', [])
            cartas_do_lote = [c for c in cartas if c['id'] in ids_lote]
            
            if busca_h:
                cartas_do_lote = [c for c in cartas_do_lote if busca_h.upper() in c['NOME'] or busca_h in c['COD_CLI']]
            
            if not cartas_do_lote and busca_h: continue

            with st.expander(f"📦 Lote {l['data']} ({len(cartas_do_lote)} itens)"):
                df_lote = pd.DataFrame(cartas_do_lote)[['NOME', 'CPF', 'VALOR', 'LOJA', 'DATA', 'MOTIVO', 'status', 'qtd_parcial']]
                st.dataframe(df_lote, use_container_width=True)
                
                out_ex = BytesIO(); df_lote.to_excel(out_ex, index=False)
                st.download_button("📥 Baixar Excel deste Lote", out_ex.getvalue(), file_name=f"Lote_{l['id']}.xlsx", key=f"dl_{l['id']}")

    # 5. RESET (SÓ ADM)
    if user_role in ["ADM", "GERENTE"]:
        with tabs[4]:
            st.warning("### ⚠️ ZONA DE PERIGO")
            st.write("Esta ação apagará todos os registros de cartas e lotes do banco de dados.")
            confirmar = st.text_input("Digite 'APAGAR' para confirmar")
            if st.button("🔥 RESETAR SISTEMA COMPLETAMENTE"):
                if confirmar == "APAGAR":
                    # Apaga cartas
                    for c in cartas:
                        fire.collection("cartas_rh").document(c['id']).delete()
                    # Apaga lotes
                    for l in lotes:
                        fire.collection("lotes_rh").document(l['id']).delete()
                    st.success("Sistema resetado com sucesso!")
                    st.rerun()
                else:
                    st.error("Palavra de confirmação incorreta.")
